"""Update resume and executive summary with AI agent swarm competencies."""
import sys
import io
import copy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

GITHUB = "github.com/rpcraighead/labinfra"

# ============================================================
# 1. UPDATE RESUME
# ============================================================
doc = Document('Ron_Craighead_AI_Cyber_Compliance.docx')

# --- Update Professional Summary ---
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'PROFESSIONAL SUMMARY':
        summary_p = doc.paragraphs[i + 1]
        summary_p.clear()
        summary_p.add_run(
            "Security-focused IT leader with 20+ years of experience implementing compliance frameworks, "
            "securing global infrastructure, and managing risk across enterprise environments. Proven expertise "
            "aligning organizations with CMMC Level 2, SOX, PCI-DSS, and GDPR requirements. Strong background "
            "in security architecture, incident response, and SecOps practices across hybrid cloud and on-premises "
            "environments. Skilled in firewall management, IDS/IPS systems, SASE/SD-WAN, virtualization, and "
            "security monitoring tools. Designed and built a multi-agent AI orchestration platform using Python, "
            "Docker, RabbitMQ, and LLM integration for autonomous infrastructure management across a Proxmox "
            "cluster with GPU-accelerated inference. Military veteran with Secret clearance history "
            "(reinvestigatable) and extensive experience supporting mission-critical, high-security systems."
        )
        break

# --- Replace Skills Table with borderless table (clean columns) ---
from docx.shared import Pt, Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

table = doc.tables[0]

# Define skill grid: 4 columns, rows
skill_grid = [
    ["CMMC Level 2 Implementation", "SOX Compliance & Audits", "PCI-DSS Compliance", "Risk Assessment"],
    ["Security Policy Development", "Incident Response", "Vendor Risk Management", "Security Architecture"],
    ["Firewall/IDS/IPS Management", "SASE/SD-WAN Security", "Identity & Access (Cisco ISE)", "SIEM (Splunk)"],
    ["Hybrid Cloud Security (AWS/Azure/OCI)", "AI Agent Orchestration", "Docker/RabbitMQ/Redis", "LLM Integration (Ollama/API)"],
    ["Proxmox Virtualization", "OpenWrt/UCI Automation", "NVIDIA GPU (A100/4070)", "Python Async/Event-Driven"],
    ["Model Context Protocol", "FastAPI Microservices", "Infrastructure as Code", "Business Continuity/DR"],
]

# Clear existing table content and resize
# Easier: remove old table, insert new borderless one at same position
parent = table._element.getparent()
table_index = list(parent).index(table._element)
parent.remove(table._element)

new_table = doc.add_table(rows=len(skill_grid), cols=4)
new_table.autofit = True

# Remove all borders (make invisible)
tbl = new_table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else etree.SubElement(tbl, qn('w:tblPr'))
borders = parse_xml(
    '<w:tblBorders %s>'
    '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    '</w:tblBorders>' % nsdecls('w')
)
tblPr.append(borders)

# Populate cells
for r_idx, row_data in enumerate(skill_grid):
    for c_idx, skill in enumerate(row_data):
        cell = new_table.cell(r_idx, c_idx)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run("\u2022 " + skill)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        # Remove cell borders too
        tcPr = cell._element.get_or_add_tcPr()
        cell_borders = parse_xml(
            '<w:tcBorders %s>'
            '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>' % nsdecls('w')
        )
        tcPr.append(cell_borders)

# Move new table to where the old one was
new_tbl_elem = new_table._tbl
new_tbl_elem.getparent().remove(new_tbl_elem)
parent.insert(table_index, new_tbl_elem)

# --- Add AI Agent achievement bullet after ML/AI bullet ---
for i, p in enumerate(doc.paragraphs):
    if 'Machine Learning/AI Simulation' in p.text:
        new_elem = copy.deepcopy(p._element)
        for child in list(new_elem):
            if child.tag.endswith('}r'):
                new_elem.remove(child)
        r = etree.SubElement(new_elem, qn('w:r'))
        t = etree.SubElement(r, qn('w:t'))
        t.text = (
            "\u2022  Designed and built CyberFlight Agent Swarm \u2014 a multi-agent AI orchestration "
            "platform with 7 specialized agents managing Proxmox VMs, Docker containers, and OpenWrt "
            "firewalls via natural language chat interface. Built with Python, FastAPI, RabbitMQ, Redis, "
            "and LLM integration (local Ollama + cloud APIs). Open source: " + GITHUB
        )
        t.set(qn('xml:space'), 'preserve')
        p._element.addnext(new_elem)
        break

# --- Add CyberFlight Lab project before Jack in the Box ---
for i, p in enumerate(doc.paragraphs):
    if 'JACK IN THE BOX' in p.text:
        jack_elem = p._element

        def make_para(text, bold=False):
            para = copy.deepcopy(jack_elem)
            for child in list(para):
                if child.tag.endswith('}r'):
                    para.remove(child)
            r = etree.SubElement(para, qn('w:r'))
            if bold:
                rPr = etree.SubElement(r, qn('w:rPr'))
                etree.SubElement(rPr, qn('w:b'))
            t = etree.SubElement(r, qn('w:t'))
            t.text = text
            t.set(qn('xml:space'), 'preserve')
            return para

        elements = [
            make_para(
                "CYBERFLIGHT LAB  |  Personal Project / Portfolio"
                "\t\t\t\t\t        01/2026 \u2013 Present",
                bold=True
            ),
            make_para("AI Infrastructure Engineer \u2014 Multi-Agent Orchestration Platform"),
            make_para(
                "\u2022  Architected a 14-container multi-agent AI system for autonomous infrastructure "
                "management across a 3-node Proxmox cluster (3 hosts, 7 VMs) with GPU passthrough (RTX 4070) "
                "for local LLM inference."
            ),
            make_para(
                "\u2022  Built 7 specialized Python/FastAPI agents: Conductor (LLM-powered orchestrator with "
                "chat UI and approval workflows), Superintendent (Proxmox VM lifecycle via API), Mercury "
                "(Docker container management), Sapper (OpenWrt firewall automation via SSH/UCI), DaVinci "
                "(IaC generation), Monitor (observability), and Scribe (documentation)."
            ),
            make_para(
                "\u2022  Implemented natural language \u2192 structured intent pipeline: user chat \u2192 LLM "
                "parsing \u2192 RabbitMQ task dispatch \u2192 agent execution \u2192 event publishing \u2192 "
                "result feedback, with dynamic infrastructure context injection and prefix-based action routing."
            ),
            make_para(
                "\u2022  Integrated multiple LLM providers (Ollama local models on RTX 4070, NVIDIA NIM API, "
                "OpenAI-compatible endpoints) with automatic failover and configurable model selection."
            ),
            make_para(
                "\u2022  Automated OpenWrt GL-MT6000 router management: firewall rules, DHCP leases, port "
                "forwards, zone forwarding, system logs, and connectivity testing \u2014 all via conversational AI."
            ),
            make_para(
                "\u2022  Infrastructure: Docker Compose, RabbitMQ (AMQP), Redis, PostgreSQL, Prometheus, "
                "Grafana, Proxmox VE API, asyncssh, aio-pika, httpx, N8N workflow automation."
            ),
            make_para(
                "\u2022  Open source: " + GITHUB
            ),
        ]

        # Insert a blank line before Jack in the Box for spacing
        blank = copy.deepcopy(jack_elem)
        for child in list(blank):
            if child.tag.endswith('}r'):
                blank.remove(child)
        jack_elem.addprevious(blank)

        # Insert elements in order (each after the previous)
        anchor = blank
        for elem in elements:
            anchor.addnext(elem)
            anchor = elem
        break

doc.save('Ron_Craighead_AI_Cyber_Compliance_v2.docx')
print("Resume saved: Ron_Craighead_AI_Cyber_Compliance_v2.docx")


# ============================================================
# 2. UPDATE EXECUTIVE SUMMARY
# ============================================================
doc2 = Document('Executive_Summary_Ron_Craighead.docx')

# --- Update KEY DIFFERENTIATORS ---
for i, p in enumerate(doc2.paragraphs):
    if 'KEY DIFFERENTIATORS:' in p.text:
        p.clear()
        p.add_run(
            "KEY DIFFERENTIATORS:\n"
            "\u2022 AIGP Certification (rare\u2014only ~5% of security professionals hold this) "
            "with hands-on AI/ML infrastructure experience\n"
            "\u2022 Built a production multi-agent AI orchestration platform "
            "(7 agents, 14 containers, natural language infrastructure management) \u2014 "
            "open source at " + GITHUB + "\n"
            "\u2022 CMMC Level 2 expertise (high demand in government contractor space)\n"
            "\u2022 Military background with reinvestigatable Secret Clearance (valuable for DoD-aligned roles)\n"
            "\u2022 Demonstrated ability to scale security at every stage: startups to Fortune 500 to public companies\n"
            "\u2022 IPO readiness experience (architected compliance infrastructure for Silvaco\u2019s successful NASDAQ IPO)\n"
            "\u2022 Cost savings track record (e.g., $250K annual savings through optimized WAN architecture)"
        )
        break

# --- Update COMPETITIVE LANDSCAPE ---
for i, p in enumerate(doc2.paragraphs):
    if 'COMPETITIVE LANDSCAPE:' in p.text:
        p.clear()
        p.add_run(
            "COMPETITIVE LANDSCAPE:\n"
            "Most security candidates in the market have either:\n"
            "\u2713 Deep compliance expertise OR deep cloud/technical expertise (Ron has both)\n"
            "\u2713 CISSP but no AI governance knowledge (Ron has AIGP\u2014differentiator)\n"
            "\u2713 No clearance background (Ron\u2019s reinvestigatable clearance is valuable)\n"
            "\u2713 No hands-on AI engineering experience (Ron built a multi-agent AI platform with "
            "LLM integration, Docker orchestration, and message-driven architecture)\n"
            "\n"
            "Conclusion: Ron has fewer direct competitors than typical security leader. His combination of "
            "compliance + AI governance + hands-on AI agent engineering + clearance background is "
            "exceptionally rare in the market."
        )
        break

# --- Update POSITIONING ---
for i, p in enumerate(doc2.paragraphs):
    if '1. POSITIONING & MESSAGING' in p.text:
        p.clear()
        p.add_run(
            "1. POSITIONING & MESSAGING\n"
            "Position Ron as \"Enterprise Security Leader with AI Engineering & Governance Expertise.\" "
            "This framing:\n"
            "\u2713 Leverages his 20+ years of credibility (CISO-track roles)\n"
            "\u2713 Differentiates him via rare AIGP certification (AI governance roles)\n"
            "\u2713 Demonstrates hands-on AI capability (built multi-agent orchestration platform)\n"
            "\u2713 Validates both traditional and forward-looking expertise\n"
            "\u2713 Justifies higher salary expectations ($180K+ range)\n"
            "\n"
            "2. TARGET PRIORITY (in order)\n"
            "Priority 1 (High conversion, fast timeline): AI Governance Manager, CMMC Lead Assessor/Manager, "
            "Security Compliance Director roles\u2014these are most actively hiring and highest demand.\n"
            "\n"
            "Priority 2 (Higher salary, longer timeline): CISO roles\u2014larger organizations, more competitive "
            "process, but higher compensation. AI engineering background strengthens candidacy.\n"
            "\n"
            "Priority 3 (Growth opportunity): AI/ML Security Engineer, Cloud Security Architect roles\u2014"
            "the CyberFlight Lab project directly demonstrates these capabilities.\n"
            "\n"
            "3. GEOGRAPHIC STRATEGY\n"
            "\u2713 San Diego local opportunities first (Insulet, SDSU, other local enterprises)\n"
            "\u2713 Remote-first government contractor roles (national market, high-paying)\n"
            "\u2713 Remote-friendly Fortune 500 companies (AWS, Salesforce, etc.)\n"
            "Rationale: Ron is San Diego-based but commands premium in remote market; "
            "local roles provide backup options."
        )
        break

# --- Update DIFFERENTIATION TACTICS ---
for i, p in enumerate(doc2.paragraphs):
    if '5. DIFFERENTIATION TACTICS' in p.text:
        p.clear()
        p.add_run(
            "5. DIFFERENTIATION TACTICS\n"
            "\u2713 Lead with AIGP + AI agent engineering in all outreach\u2014this is a market edge\n"
            "\u2713 Reference CyberFlight Agent Swarm project (" + GITHUB + ") "
            "as proof of hands-on AI capability\n"
            "\u2713 Reference IPO/scaling experience in CISO conversations\n"
            "\u2713 Reference clearance background in DoD/government contractor conversations\n"
            "\u2713 Emphasize compliance-to-innovation balance (security enables business, not constraint)\n"
            "\u2713 Highlight multi-agent architecture as proof of systems thinking at scale"
        )
        break

# --- Update DIFFERENTIATION conclusion ---
for i, p in enumerate(doc2.paragraphs):
    if '2. DIFFERENTIATION: The combination of CISSP' in p.text:
        p.clear()
        p.add_run(
            "2. DIFFERENTIATION: The combination of CISSP + AIGP + CMMC expertise + military clearance "
            "background + hands-on AI agent engineering is exceptionally rare and valuable. Ron should lead "
            "with this combination in all outreach. The CyberFlight Agent Swarm project (" + GITHUB + ") "
            "provides tangible proof of AI engineering capability that most security leaders cannot demonstrate."
        )
        break

# --- Update AI Governance row in target roles table ---
table2 = doc2.tables[0]
for row in table2.rows:
    if 'AI Governance' in row.cells[0].text:
        row.cells[3].paragraphs[0].clear()
        row.cells[3].paragraphs[0].add_run(
            "AIGP certification (rare), hands-on AI agent platform "
            "(" + GITHUB + "), compliance background"
        )
        break

doc2.save('Executive_Summary_Ron_Craighead_v2.docx')
print("Executive Summary saved: Executive_Summary_Ron_Craighead_v2.docx")
