"""Convert LAB.md to a formatted Word document."""
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

def parse_markdown(md_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        return f.read()

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear',
    })
    shading.append(shading_elm)

def add_formatted_run(paragraph, text):
    """Parse **bold** and `code` inline formatting."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        else:
            paragraph.add_run(part)

def build_docx(md_text, output_path, diagram_path=None):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(2)

    # Heading styles
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.color.rgb = RGBColor(0x0A, 0x1A, 0x3A)
        if level == 1:
            hs.font.size = Pt(22)
            hs.paragraph_format.space_before = Pt(0)
            hs.paragraph_format.space_after = Pt(12)
        elif level == 2:
            hs.font.size = Pt(16)
            hs.paragraph_format.space_before = Pt(18)
            hs.paragraph_format.space_after = Pt(8)
        elif level == 3:
            hs.font.size = Pt(13)
            hs.paragraph_format.space_before = Pt(14)
            hs.paragraph_format.space_after = Pt(6)
        elif level == 4:
            hs.font.size = Pt(11)
            hs.paragraph_format.space_before = Pt(10)
            hs.paragraph_format.space_after = Pt(4)

    lines = md_text.split('\n')
    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i]

        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
                # Add light gray background via shading
                pPr = p._element.get_or_add_pPr()
                shd = pPr.makeelement(qn('w:shd'), {
                    qn('w:fill'): 'F2F2F2',
                    qn('w:val'): 'clear',
                })
                pPr.append(shd)
                in_code_block = False
                code_lines = []
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Empty lines
        if not line.strip():
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Clean markdown formatting from heading
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            p = doc.add_heading(text, level=level)
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a border line
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): 'CCCCCC',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Tables
        if '|' in line and i + 1 < len(lines) and re.match(r'\s*\|[\s\-|]+\|\s*$', lines[i + 1]):
            # Parse table
            table_lines = []
            table_lines.append(line)
            i += 1  # skip header
            i += 1  # skip separator
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1

            # Parse cells
            rows = []
            for tl in table_lines:
                cells = [c.strip() for c in tl.strip().strip('|').split('|')]
                rows.append(cells)

            if not rows:
                continue

            num_cols = len(rows[0])
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            table.style = 'Table Grid'

            for r_idx, row in enumerate(rows):
                for c_idx, cell_text in enumerate(row):
                    if c_idx >= num_cols:
                        break
                    cell = table.cell(r_idx, c_idx)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    add_formatted_run(p, cell_text)

                    # Header row styling
                    if r_idx == 0:
                        set_cell_shading(cell, '0A1A3A')
                        for run in p.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.size = Pt(10)
                    else:
                        # Alternate row shading
                        if r_idx % 2 == 0:
                            set_cell_shading(cell, 'F5F5F5')
                        for run in p.runs:
                            run.font.size = Pt(10)
            continue

        # Bullet points
        bullet_match = re.match(r'^(\s*)- (.*)', line)
        if bullet_match:
            indent_level = len(bullet_match.group(1)) // 2
            text = bullet_match.group(2)
            p = doc.add_paragraph(style='List Bullet')
            p.clear()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(1.2 * indent_level)
            add_formatted_run(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_run(p, line)
        i += 1

    # Insert diagram image after "Network Map" heading if available
    # (We'll skip this since the image is already on the website)

    doc.save(output_path)
    print(f"Saved: {output_path}")

if __name__ == '__main__':
    script_dir = os.path.dirname(os.path.abspath(__file__))
    lab_md = os.path.join(script_dir, '..', 'LAB.md')
    output = os.path.join(script_dir, '..', 'CyberFlight_Lab_Documentation.docx')
    md_text = parse_markdown(lab_md)
    build_docx(md_text, output)
